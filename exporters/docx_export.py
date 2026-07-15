"""
exporters/docx_export.py
---------------------------
Word (.docx) newsletter exporter using python-docx.
"""

from __future__ import annotations
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_docx(newsletter: Dict[str, Any], out_path: str) -> str:
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("FMCG Deal Intelligence Newsletter", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Generated {newsletter.get('generated_at', '')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].italic = True

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(newsletter.get("executive_summary", ""))

    doc.add_heading("Top Deals", level=1)
    deals = newsletter.get("deals", [])
    if deals:
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        headers = ["Title", "Published", "Deal Type", "Amount", "Sector", "Source", "Credibility"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for d in deals:
            row = table.add_row().cells
            row[0].text = d.get("title", "")[:120]
            row[1].text = d.get("published", "Date unknown")
            row[2].text = d.get("deal_type", "").replace("_", " ").title()
            row[3].text = str(d.get("amount") or "—")
            row[4].text = d.get("sector", "FMCG")
            row[5].text = d.get("source", "")
            row[6].text = d.get("credibility_label", "")
    else:
        doc.add_paragraph("No deals met the relevance threshold in this run.")

    doc.add_heading("Market Trend Snapshot", level=1)
    trend = newsletter.get("trend_snapshot", {})
    for dtype, count in sorted(trend.items(), key=lambda x: -x[1]):
        doc.add_paragraph(f"{dtype.replace('_', ' ').title()}: {count} deal(s)", style="List Bullet")

    doc.add_heading("Methodology & Assumptions", level=1)
    for line in [
        "Sources: free public RSS feeds (Google News, Reuters, ET, Mint, Business Standard, CNBC, Financial Express, Yahoo Finance).",
        "Relevance: rule-based keyword scoring (FMCG company + deal-action co-occurrence).",
        "Credibility: static publisher-tier table — a proxy for editorial track record, not a fact-check of the specific claim.",
        "Deduplication: exact URL/title match, then TF-IDF cosine similarity on cleaned article text.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(out_path)
    return out_path
